import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path):
    """Extract text from .pdf, .docx, or .doc files."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text_chunks = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_chunks.append(page_text)
        return "\n".join(text_chunks).strip()

    if ext == ".docx":
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()]).strip()

    if ext == ".doc":
        try:
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()]).strip()
        except Exception as e:
            raise ValueError(f"Unsupported or corrupt .doc file: {os.path.basename(file_path)}") from e

    raise ValueError(f"Unsupported file format: {ext}")
