import os
import fitz  # PyMuPDF
from docx import Document
# Removed 'import text_extract as te' as it was a circular import

def extract_text_from_pdf(file_path):
    """Extracts text from .pdf, .docx, or .doc files."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        print(f"Extracting text from PDF: {file_path}")
        with fitz.open(file_path) as pdf:
            text = ""
            for page in pdf:
                text += page.get_text()
            return text.strip()
        
    elif ext == '.docx':
        print(f"Extracting text from DOCX: {file_path}")
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        return text.strip()
    
    elif ext == '.doc':
        # Basic .doc support (might be less reliable than .docx)
        # For full .doc support, antiword (Linux) or pywin32 (Windows) is often needed
        # This implementation with 'docx' might handle some .doc files, but it's not guaranteed.
        # We'll try to open it as a .docx, which sometimes works for simple .doc files.
        try:
            print(f"Attempting to extract text from DOC: {file_path}")
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            return text.strip()
        except Exception as e:
            print(f"Could not read .doc file {file_path} directly: {e}")
            raise ValueError(f"Unsupported or corrupt .doc file: {os.path.basename(file_path)}")
    
    else:
        raise ValueError("Unsupported file format: {}".format(ext))
